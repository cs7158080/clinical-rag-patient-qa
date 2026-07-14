"""
adapter_a.py — Family A parser for diagnosis and clinic_visit_summary documents.

Parses .docx files with a fixed section structure and extracts:
- Header fields (name, date, date_of_birth, national_id, hmo_name)
- Section text (one entry per fixed section key)
- Domain findings (diagnosis only, from the ממצאי האבחון section)
"""

import logging
import re
from typing import Optional

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Section and domain constants
# ---------------------------------------------------------------------------

# Exact Hebrew header text as it appears in the .docx paragraphs
SECTION_HEADERS = [
    "רקע",
    "מהלך האבחון",
    "ממצאי האבחון",
    "סיכום והמלצות",
    "תמצית אבחון",
]

# Normalised section keys (used as 'section' field in SQLite)
SECTION_KEYS = [
    "רקע",
    "מהלך_האבחון",
    "ממצאי_האבחון",
    "סיכום_והמלצות",
    "תמצית_אבחון",
]

# Treatment-worded headings used by generated סיכום טיפול files — same section keys
TREATMENT_SECTION_HEADERS: dict[str, str] = {
    "מהלך הטיפול": "מהלך_האבחון",
    "ממצאי הטיפול": "ממצאי_האבחון",
}

# Mapping from header text → section key (recognises both heading sets)
HEADER_TO_KEY: dict[str, str] = dict(zip(SECTION_HEADERS, SECTION_KEYS)) | TREATMENT_SECTION_HEADERS

# 13 domain sub-headers inside ממצאי האבחון (diagnosis only)
DOMAIN_HEADERS = [
    "פרגמטיקה ותקשורת",
    "הבנת שפה",
    "הבעת שפה",
    "לקסיקון",
    "תחביר",
    "מורפולוגיה",
    "התארגנות להבעת מלל מורכב",
    "מובנות הדיבור",
    "מודעות פונולוגית",
    "זיכרון שמיעתי",
    "אורל מוטור",
    "אכילה",
    "שטף",
]

# Map sub-domains to their parent category
PARENT_DOMAINS: dict[str, str] = {
    "לקסיקון": "הבעת שפה",
    "תחביר": "הבעת שפה",
    "מורפולוגיה": "הבעת שפה",
    "התארגנות להבעת מלל מורכב": "הבעת שפה",
}

# Parent-only categories: headers above their child domains — never stored
# as findings rows of their own. Derived from PARENT_DOMAINS.
PARENT_ONLY_DOMAINS: set[str] = set(PARENT_DOMAINS.values())

# Signature-block sentinel — clinician licence number indicator
_SIGNATURE_SENTINEL = "מ.ר."


# ---------------------------------------------------------------------------
# Template type detection
# ---------------------------------------------------------------------------

def detect_template_type(filename: str) -> Optional[str]:
    """Detect the Family A template type from the document filename.

    Parameters
    ----------
    filename:
        The basename of the .docx file (with or without extension).

    Returns
    -------
    'diagnosis' if the filename starts with 'סיכום אבחון',
    'clinic_visit_summary' if it starts with 'סיכום טיפול',
    None otherwise.
    """
    if filename.startswith("סיכום אבחון"):
        return "diagnosis"
    if filename.startswith("סיכום טיפול"):
        return "clinic_visit_summary"
    return None


# ---------------------------------------------------------------------------
# Header field extraction
# ---------------------------------------------------------------------------

def extract_header_fields(doc) -> dict:
    """Scan the first ~10 paragraphs of *doc* for known header field patterns.

    Fields extracted
    ----------------
    name           : text after 'שם:' or 'שם '
    date           : text after 'תאריך:'
    date_of_birth  : text after 'ת.ל.:'
    national_id    : text after 'ת.ז.:'
    hmo_name       : text after 'קופת חולים:' / 'קופת חולים ' / 'HMO:'

    Parameters
    ----------
    doc:
        A python-docx Document object.

    Returns
    -------
    Dict with keys: name, date, date_of_birth, national_id, hmo_name.
    All values are str or None.
    """
    result: dict[str, Optional[str]] = {
        "name": None,
        "date": None,
        "date_of_birth": None,
        "national_id": None,
        "hmo_name": None,
    }

    # Patterns: each tuple is (field_key, compiled_regex)
    # Each regex captures the value in group 1.
    patterns = [
        ("name",          re.compile(r"שם\s*[:\s]\s*(.+)")),
        ("date",          re.compile(r"תאריך\s*:\s*(.+)")),
        ("date_of_birth", re.compile(r"ת\.ל\.\s*:\s*(.+)")),
        ("national_id",   re.compile(r"ת\.ז\.\s*:\s*(.+)")),
        ("hmo_name",      re.compile(r"(?:קופת חולים|HMO)\s*[:\s]\s*(.+)", re.IGNORECASE)),
    ]

    paragraphs = doc.paragraphs[:10]
    for para in paragraphs:
        text = para.text.strip()
        if not text:
            continue
        for field_key, pattern in patterns:
            if result[field_key] is not None:
                # Already found this field
                continue
            match = pattern.search(text)
            if match:
                value = match.group(1).strip()
                if value:
                    result[field_key] = value

    return result


# ---------------------------------------------------------------------------
# Main parser
# ---------------------------------------------------------------------------

def parse_family_a(
    file_path: str,
    patient_folder_name: str,
    template_type: str,
) -> dict:
    """Parse a Family A .docx document and return structured data.

    Parameters
    ----------
    file_path:
        Absolute path to the .docx file.
    patient_folder_name:
        The name of the patient's folder (used as the canonical patient
        identifier before de-identification).
    template_type:
        'diagnosis' or 'clinic_visit_summary'.

    Returns
    -------
    {
        "template_type": str,
        "patient_folder_name": str,
        "header": dict,           # from extract_header_fields
        "sections": dict,         # section_key -> concatenated paragraph text
        "domains": dict,          # domain_name -> text (diagnosis only, else {})
        "file_path": str,
    }
    """
    from docx import Document  # type: ignore

    try:
        doc = Document(file_path)
    except Exception as exc:
        logger.error("parse_family_a: failed to open document %s: %s", file_path, exc)
        raise

    header = extract_header_fields(doc)

    # Initialise all sections with empty strings (missing sections are kept as "")
    sections: dict[str, str] = {key: "" for key in SECTION_KEYS}
    domains: dict[str, str] = {}

    current_section_key: Optional[str] = None
    current_domain: Optional[str] = None
    section_lines: dict[str, list[str]] = {key: [] for key in SECTION_KEYS}
    domain_lines: dict[str, list[str]] = {}
    blocked: bool = False  # True once signature block sentinel is encountered

    for para in doc.paragraphs:
        raw_text = para.text
        stripped = raw_text.strip()

        # Empty paragraph — skip (but reset blocked flag per section only)
        if not stripped:
            continue

        # --- Signature block sentinel ---
        # If the clinician licence marker or the closing salutation appears,
        # stop collecting for the current section.
        if _SIGNATURE_SENTINEL in stripped or stripped.startswith("בברכה"):
            blocked = True
            current_section_key = None
            current_domain = None
            continue

        # --- Section header detection ---
        # Use startswith() to handle colons and whitespace variants
        section_key = None
        for section_header_text, key in HEADER_TO_KEY.items():
            if stripped.rstrip(':').startswith(section_header_text.rstrip(':')):
                section_key = key
                break

        if section_key is not None:
            current_section_key = section_key
            current_domain = None
            blocked = False  # new section resets the blocked flag
            continue

        # --- Domain header detection (inside ממצאי_האבחון) ---
        if current_section_key == "ממצאי_האבחון":
            matched_domain = None

            for domain in DOMAIN_HEADERS:
                pattern = rf"^{re.escape(domain)}\s*[:\-–]?\s*(.*)$"
                match = re.match(pattern, stripped)

                if match:
                    matched_domain = domain
                    current_domain = domain
                    domain_lines.setdefault(domain, [])

                    inline_text = match.group(1).strip()
                    if inline_text:
                        domain_lines[domain].append(inline_text)
                    break

            if matched_domain is not None:
                continue

        # --- Content collection ---
        if blocked:
            continue

        if current_section_key is not None:
            if current_domain is not None:
                # Collect into the current domain
                domain_lines[current_domain].append(stripped)
            else:
                # Collect into the current section
                section_lines[current_section_key].append(stripped)

    # Assemble section texts
    for key in SECTION_KEYS:
        section_texts = section_lines.get(key, [])
        sections[key] = "\n".join(section_texts).strip()

    # Assemble domain texts (both Family A types).
    # Parent-only categories are section headers, not domains — skipped entirely.
    for domain_name, lines in domain_lines.items():
        if domain_name in PARENT_ONLY_DOMAINS:
            continue
        domains[domain_name] = "\n".join(lines).strip()

    logger.info(
        "parse_family_a: parsed %s (%s) — sections=%d domains=%d",
        file_path,
        template_type,
        sum(1 for v in sections.values() if v),
        len(domains),
    )

    return {
        "template_type": template_type,
        "patient_folder_name": patient_folder_name,
        "header": header,
        "sections": sections,
        "domains": domains,
        "parent_domains": PARENT_DOMAINS,
        "file_path": file_path,
    }
