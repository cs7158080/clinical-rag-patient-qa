"""
adapter_b.py — Family B parser for treatment_plan documents.

Parses .docx files that contain:
  Layer 1 — A goals table (doc.tables[0]): rows of (session_date, goals_text)
  Layer 2 — Free-text session summaries appended below the table(s),
             delimited by date lines.

Date formats supported (in priority order):
  DD/MM/YYYY, D/M/YYYY, DD.MM.YYYY, D.M.YYYY
All dates are normalised to ISO 8601 YYYY-MM-DD before storage.
"""

import logging
import re
from datetime import datetime, date
from typing import Optional

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Date format registry (tried in order)
# ---------------------------------------------------------------------------

# Each entry: (strptime_format, compiled_regex_for_quick_pre-check)
LAYER2_DATE_FORMATS = [
    ("%d/%m/%Y", re.compile(r"^\d{2}/\d{2}/\d{4}$")),
    ("%d/%m/%Y", re.compile(r"^\d{1,2}/\d{1,2}/\d{4}$")),
    ("%d.%m.%Y", re.compile(r"^\d{2}\.\d{2}\.\d{4}$")),
    ("%d.%m.%Y", re.compile(r"^\d{1,2}\.\d{1,2}\.\d{4}$")),
]


# ---------------------------------------------------------------------------
# Date parsing helpers
# ---------------------------------------------------------------------------

def parse_layer2_date(line: str) -> Optional[str]:
    """Try to parse *line* as a session-boundary date.

    Formats are tried in the order defined by LAYER2_DATE_FORMATS.
    The first successful parse wins and is returned as an ISO 8601 string
    (YYYY-MM-DD).  Returns None if no format matches.

    Parameters
    ----------
    line:
        A single stripped line of text from the document.

    Returns
    -------
    ISO date string ('YYYY-MM-DD') or None.
    """
    stripped = line.strip()
    for fmt, pattern in LAYER2_DATE_FORMATS:
        if pattern.match(stripped):
            try:
                parsed: date = datetime.strptime(stripped, fmt).date()
                return parsed.isoformat()
            except ValueError:
                # strptime failed despite regex match (e.g. month=13) — try next
                continue
    return None


def _normalize_table_date(raw: str) -> Optional[str]:
    """Attempt to parse a table cell value as a date and return ISO 8601.

    Uses the same format list as parse_layer2_date.  Returns the original
    string unchanged if no format matches (best-effort — we do not block on
    unparseable table dates).
    """
    if not raw or not raw.strip():
        return None
    iso = parse_layer2_date(raw.strip())
    if iso:
        return iso
    # Return raw value as fallback — caller handles None from table cell
    return raw.strip() or None


# ---------------------------------------------------------------------------
# Header extraction
# ---------------------------------------------------------------------------

def extract_treatment_header(doc) -> dict:
    """Extract file-level header fields from the first paragraphs of *doc*.

    Fields
    ------
    name           : text after 'שם:' / 'שם '
    date_of_birth  : text after 'ת.ל.:'
    start_date     : text after 'תאריך התחלת הטיפול:' (ISO if parseable, else raw)
    main_goal      : text after 'מטרת על:'
    notes          : text after 'הערות:'

    Parameters
    ----------
    doc:
        A python-docx Document object.

    Returns
    -------
    Dict with the five keys above.  All values are str or None.
    """
    result: dict[str, Optional[str]] = {
        "name": None,
        "date_of_birth": None,
        "start_date": None,
        "main_goal": None,
        "notes": None,
    }

    patterns = [
        ("name",          re.compile(r"שם\s*[:\s]\s*(.+)")),
        ("date_of_birth", re.compile(r"ת\.ל\.\s*:\s*(.+)")),
        ("start_date",    re.compile(r"תאריך התחלת הטיפול\s*:\s*(.+)")),
        ("main_goal",     re.compile(r"מטרת על\s*:\s*(.+)")),
        ("notes",         re.compile(r"הערות\s*:\s*(.+)")),
    ]

    for para in doc.paragraphs[:15]:
        text = para.text.strip()
        if not text:
            continue
        for field_key, pattern in patterns:
            if result[field_key] is not None:
                continue
            match = pattern.search(text)
            if match:
                value = match.group(1).strip()
                if value:
                    result[field_key] = value

    # Normalise start_date to ISO if possible
    if result["start_date"]:
        iso = parse_layer2_date(result["start_date"])
        if iso:
            result["start_date"] = iso

    return result


# ---------------------------------------------------------------------------
# Main parser
# ---------------------------------------------------------------------------

def parse_treatment_plan(file_path: str, patient_folder_name: str) -> dict:
    """Parse a Family B (treatment_plan) .docx document.

    Parameters
    ----------
    file_path:
        Absolute path to the .docx file.
    patient_folder_name:
        The name of the patient's folder.

    Returns
    -------
    {
        "template_type": "treatment_plan",
        "patient_folder_name": str,
        "header": dict,
        "goals_rows": list of {"session_date": str, "goals_text": str},
        "session_blocks": list of {"session_date": str, "session_text": str},
        "file_path": str,
    }
    """
    from docx import Document  # type: ignore

    try:
        doc = Document(file_path)
    except Exception as exc:
        logger.error("parse_treatment_plan: failed to open document %s: %s", file_path, exc)
        raise

    header = extract_treatment_header(doc)

    # ------------------------------------------------------------------
    # Layer 1 — Goals table (doc.tables[0])
    # ------------------------------------------------------------------
    goals_rows: list[dict] = []

    if doc.tables:
        table = doc.tables[0]
        for row in table.rows:
            cells = row.cells
            if len(cells) < 2:
                continue
            raw_date = cells[0].text.strip()
            goals_text = cells[1].text.strip()

            if not raw_date and not goals_text:
                # Blank row — skip
                continue

            iso_date = _normalize_table_date(raw_date)
            if iso_date and goals_text:
                goals_rows.append({
                    "session_date": iso_date,
                    "goals_text": goals_text,
                })
            elif goals_text:
                # Date missing or unparseable — store with empty date
                logger.warning(
                    "parse_treatment_plan: goals table row has no parseable date "
                    "(raw=%r) in %s — skipping row",
                    raw_date,
                    file_path,
                )

    # ------------------------------------------------------------------
    # Layer 2 — Session summaries (free text paragraphs after table(s))
    # ------------------------------------------------------------------
    # We scan ALL paragraphs in the document; paragraphs that are inside
    # a table body are skipped by python-docx's doc.paragraphs (it returns
    # only body-level paragraphs, not table-cell paragraphs).
    # So all paragraphs here are candidates for Layer 2 content.

    session_blocks: list[dict] = []
    current_date: Optional[str] = None
    current_lines: list[str] = []
    found_any_date = False

    for para in doc.paragraphs:
        line = para.text.strip()
        if not line:
            continue

        iso = parse_layer2_date(line)
        if iso:
            # Flush previous block
            if current_date is not None and current_lines:
                session_blocks.append({
                    "session_date": current_date,
                    "session_text": "\n".join(current_lines).strip(),
                })
            current_date = iso
            current_lines = []
            found_any_date = True
        else:
            if current_date is not None:
                current_lines.append(line)
            # Lines before the first date marker are ignored for Layer 2
            # (they are typically header material already captured above)

    # Flush the last open block
    if current_date is not None and current_lines:
        session_blocks.append({
            "session_date": current_date,
            "session_text": "\n".join(current_lines).strip(),
        })

    # Edge case: no date markers found anywhere in the free-text area
    if not found_any_date and doc.paragraphs:
        fallback_date = header.get("start_date") or datetime.today().date().isoformat()
        logger.warning(
            "parse_treatment_plan: no session-date markers found in Layer 2 of %s. "
            "Treating all text as a single block with session_date=%s",
            file_path,
            fallback_date,
        )
        all_lines = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        if all_lines:
            session_blocks.append({
                "session_date": fallback_date,
                "session_text": "\n".join(all_lines).strip(),
            })

    logger.info(
        "parse_treatment_plan: parsed %s — goals_rows=%d session_blocks=%d",
        file_path,
        len(goals_rows),
        len(session_blocks),
    )

    return {
        "template_type": "treatment_plan",
        "patient_folder_name": patient_folder_name,
        "header": header,
        "goals_rows": goals_rows,
        "session_blocks": session_blocks,
        "file_path": file_path,
    }
