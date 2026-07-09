# LlamaIndex Generate Summary Workflow.
from llama_index.core.workflow import Workflow, StartEvent, StopEvent, step, Context, Event
from llama_index.llms.anthropic import Anthropic as LlamaAnthropic
from typing import Union
import os
import io
import json
import html
import hashlib
import logging
from docx import Document as DocxDocument
from bidi.algorithm import get_display

from app.storage import db
from app.storage.models import FamilyAChunk
from app.deidentification.reid_map import reidentify_text, load as load_reid_map, reverse_lookup
from app.prompts.generation import GENERATE_SUMMARY_PROMPT, parse_generated_sections, SECTION_KEYS
from app.prompts.qa import NO_SESSION_MESSAGE, ERROR_MESSAGE
from app.config import AppConfig

logger = logging.getLogger(__name__)

SECTION_LABELS = {
    'רקע': 'רקע:',
    'מהלך_האבחון': 'מהלך האבחון:',
    'ממצאי_האבחון': 'ממצאי האבחון:',
    'סיכום_והמלצות': 'סיכום והמלצות:',
    'תמצית_אבחון': 'תמצית אבחון',
}

NO_BASE_DOCUMENT_MESSAGE = "לא נמצא מסמך בסיס (אבחון או סיכום ביקור קודם) — לא ניתן לייצר סיכום"


def _rtl(text: str) -> str:
    return html.escape(get_display(str(text))) if text else ''


def _save_pdf(docx_path: str, patient_name: str, national_id: str, date_of_birth: str,
              session_date: str, sections: dict) -> str:
    from reportlab.lib.pagesizes import A4
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.enums import TA_RIGHT
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    pdf_path = docx_path.replace('.docx', '.pdf')

    font_name = 'Helvetica'
    for candidate in [r'C:\Windows\Fonts\arial.ttf', r'C:\Windows\Fonts\ARIAL.TTF']:
        if os.path.exists(candidate):
            try:
                pdfmetrics.registerFont(TTFont('ArialHeb', candidate))
                font_name = 'ArialHeb'
            except Exception:
                pass
            break

    meta_style = ParagraphStyle('Meta', fontName=font_name, fontSize=11,
                                alignment=TA_RIGHT, leading=16, spaceAfter=2)
    heading_style = ParagraphStyle('SectionHeading', fontName=font_name, fontSize=13,
                                   alignment=TA_RIGHT, leading=18, spaceBefore=10, spaceAfter=4)
    body_style = ParagraphStyle('Body', fontName=font_name, fontSize=11,
                                alignment=TA_RIGHT, leading=16, spaceAfter=2)

    doc = SimpleDocTemplate(
        pdf_path, pagesize=A4,
        rightMargin=50, leftMargin=50, topMargin=50, bottomMargin=50,
    )

    story = []
    story.append(Paragraph(_rtl(f'שם: {patient_name}'), meta_style))
    story.append(Paragraph(_rtl(f'ת.ז.: {national_id}'), meta_style))
    story.append(Paragraph(_rtl(f'ת.ל.: {date_of_birth}'), meta_style))
    story.append(Paragraph(_rtl(f'תאריך: {session_date}'), meta_style))
    story.append(Spacer(1, 14))

    for key, label in SECTION_LABELS.items():
        story.append(Paragraph(_rtl(label), heading_style))
        for line in (sections.get(key) or '').split('\n'):
            story.append(Paragraph(_rtl(line), body_style))
        story.append(Spacer(1, 6))

    doc.build(story)
    return pdf_path


# ---------------------------------------------------------------------------
# Events
# ---------------------------------------------------------------------------

class GenerateStartEvent(StartEvent):
    patient_id: str
    date_from: str   # first session in range (ISO 8601)
    date_to: str     # last session in range (ISO 8601)


class SourcesEvent(Event):
    sessions_text: str       # all sessions in range, joined by \n---\n
    goals_text: str          # all goals rows in range, joined by \n---\n
    base_document: dict      # {section_key: text} from latest clinic_visit_summary or diagnosis
    patient_id: str
    date_from: str
    date_to: str


class TokenizedDocEvent(Event):
    sections: dict
    patient_id: str
    date_to: str   # used as session_date when writing to SQLite


class ReidentifiedDocEvent(Event):
    sections: dict             # re-identified — for the .docx body only
    tokenized_sections: dict   # de-identified — the ONLY version stored in SQLite
    patient_id: str
    date_to: str
    patient_name: str
    national_id: str
    date_of_birth: str


# ---------------------------------------------------------------------------
# Workflow
# ---------------------------------------------------------------------------

class GenerateSummaryWorkflow(Workflow):
    def __init__(self, config: AppConfig, db_path: str, reid_map_path: str, **kwargs):
        super().__init__(**kwargs)
        self._config = config
        self._db_path = db_path
        self._reid_map_path = reid_map_path
        self._llm = LlamaAnthropic(
            model=config.anthropic.generation_model,
            api_key=config.anthropic_api_key,
            temperature=config.anthropic.temperature_generation,
            max_tokens=10000
        )

    @step
    async def fetch_sources_step(
        self, ctx: Context, ev: GenerateStartEvent
    ) -> Union[SourcesEvent, StopEvent]:
        # Sessions in range
        sessions = db.fetch_treatment_sessions(
            self._db_path, ev.patient_id, ev.date_from, ev.date_to
        )
        if not sessions:
            return StopEvent(result=NO_SESSION_MESSAGE)

        sessions_text = '\n---\n'.join(s.session_text_deidentified for s in sessions)

        # Goals in range
        goals_rows = db.fetch_treatment_goals_in_range(
            self._db_path, ev.patient_id, ev.date_from, ev.date_to
        )
        goals_text = '\n---\n'.join(g.goals_text_deidentified for g in goals_rows) if goals_rows else ''

        # Base document: latest clinic_visit_summary → latest diagnosis → error
        base_document = db.fetch_latest_family_a_as_dict(
            self._db_path, ev.patient_id, 'clinic_visit_summary'
        )
        if base_document is None:
            logger.info(
                "No clinic_visit_summary found for patient=%s, falling back to diagnosis",
                ev.patient_id,
            )
            base_document = db.fetch_latest_family_a_as_dict(
                self._db_path, ev.patient_id, 'diagnosis'
            )
        if base_document is None:
            logger.error("No base document found for patient=%s", ev.patient_id)
            return StopEvent(result=NO_BASE_DOCUMENT_MESSAGE)

        return SourcesEvent(
            sessions_text=sessions_text,
            goals_text=goals_text,
            base_document=base_document,
            patient_id=ev.patient_id,
            date_from=ev.date_from,
            date_to=ev.date_to,
        )

    @step
    async def generate_summary_step(
        self, ctx: Context, ev: SourcesEvent
    ) -> Union[TokenizedDocEvent, StopEvent]:
        prompt = GENERATE_SUMMARY_PROMPT.format(
            base_document_json=json.dumps(ev.base_document, ensure_ascii=False, indent=2),
            date_from=ev.date_from,
            date_to=ev.date_to,
            sessions_text=ev.sessions_text,
            goals_text=ev.goals_text or 'אין מטרות',
        )
        try:
            logger.info(
                "Generating summary for patient=%s range=%s to %s",
                ev.patient_id, ev.date_from, ev.date_to,
            )
            response = await self._llm.acomplete(prompt)
            raw = response.text.strip()
            if not raw:
                return StopEvent(result=ERROR_MESSAGE)
            sections = parse_generated_sections(raw)
            return TokenizedDocEvent(
                sections=sections,
                patient_id=ev.patient_id,
                date_to=ev.date_to,
            )
        except Exception as e:
            logger.error("Summary generation error: %s", e)
            return StopEvent(result=ERROR_MESSAGE)

    @step
    async def build_doc_step(
        self, ctx: Context, ev: TokenizedDocEvent
    ) -> Union[ReidentifiedDocEvent, StopEvent]:
        try:
            reid_map = load_reid_map(self._reid_map_path)
        except Exception:
            return StopEvent(result="לא ניתן לייצר סיכום — מפת הזיהוי חסרה")

        reidentified_sections = {k: reidentify_text(reid_map, v) for k, v in ev.sections.items()}

        metadata = db.get_patient_metadata(self._db_path, ev.patient_id)
        patient_token = f"PERSON_{ev.patient_id}"
        patient_name = reverse_lookup(reid_map, patient_token) or ev.patient_id
        national_id = metadata.get('national_id', '')
        date_of_birth = metadata.get('date_of_birth', '')

        return ReidentifiedDocEvent(
            sections=reidentified_sections,
            tokenized_sections=ev.sections,
            patient_id=ev.patient_id,
            date_to=ev.date_to,
            patient_name=patient_name,
            national_id=national_id,
            date_of_birth=date_of_birth,
        )

    @step
    async def save_doc_step(self, ctx: Context, ev: ReidentifiedDocEvent) -> StopEvent:
        reid_map = load_reid_map(self._reid_map_path)
        patient_folder_name = reverse_lookup(reid_map, f"PERSON_{ev.patient_id}") or ev.patient_id
        patients_root = self._config.patients_root
        patient_folder = os.path.join(patients_root, patient_folder_name)

        existing = [f for f in os.listdir(patient_folder) if f.startswith('סיכום ביקור')]
        next_num = len(existing) + 1
        base_name = f"סיכום ביקור {next_num} {patient_folder_name}"
        docx_path = os.path.join(patient_folder, base_name + '.docx')

        doc = DocxDocument()
        doc.add_paragraph(f"שם: {ev.patient_name}")
        doc.add_paragraph(f"ת.ז.: {ev.national_id}")
        doc.add_paragraph(f"ת.ל.: {ev.date_of_birth}")
        doc.add_paragraph(f"תאריך: {ev.date_to}")

        for key, label in SECTION_LABELS.items():
            doc.add_paragraph(label, style='Heading 2')
            doc.add_paragraph(ev.sections.get(key, ''))

        try:
            doc.save(docx_path)
        except Exception as e:
            logger.error("Failed to save summary doc: %s", e)
            return StopEvent(result=f"שגיאה בשמירת הקובץ: {e}")

        buf = io.BytesIO()
        doc.save(buf)
        file_hash = hashlib.sha256(buf.getvalue()).hexdigest()

        # Write de-identified (tokenized) sections to SQLite — never the
        # re-identified text; session_date = date_to (last session in range)
        for section_key, text in ev.tokenized_sections.items():
            chunk = FamilyAChunk(
                patient_id=ev.patient_id,
                template_type='clinic_visit_summary',
                session_date=ev.date_to,
                section=section_key,
                text_deidentified=text,
                source_file_path=docx_path,
            )
            db.insert_family_a_chunk(self._db_path, chunk)

        db.mark_file_ingested(self._db_path, docx_path, file_hash)

        pdf_path = None
        try:
            pdf_path = _save_pdf(
                docx_path,
                patient_name=ev.patient_name,
                national_id=ev.national_id,
                date_of_birth=ev.date_of_birth,
                session_date=ev.date_to,
                sections=ev.sections,
            )
            logger.info("PDF saved: %s", pdf_path)
        except Exception as e:
            logger.error("PDF generation failed: %s", e)

        logger.info("Summary generated and saved: %s", docx_path)
        return StopEvent(result=json.dumps({"docx": docx_path, "pdf": pdf_path}))


async def run_generate_summary(
    patient_id: str,
    date_from: str,
    date_to: str,
    config: AppConfig,
    db_path: str,
    reid_map_path: str,
) -> str:
    workflow = GenerateSummaryWorkflow(
        config=config,
        db_path=db_path,
        reid_map_path=reid_map_path,
        timeout=120,
    )
    result = await workflow.run(
        patient_id=patient_id,
        date_from=date_from,
        date_to=date_to,
    )
    return result
