"""
docx_builder.py — Build the generated summary .docx from a physical skeleton file.

The skeleton is an existing .docx opened as-is (the patient's previous summary /
diagnosis, or the bundled clinic template as fallback). Formatting is never rebuilt:
untouched paragraphs stay exactly as in the skeleton, and injected paragraphs clone
the paragraph and run properties (pPr / rPr) of the content they replace, so font,
size, RTL direction and indentation are inherited.
"""

import copy
import logging
from typing import Optional

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

from app.ingestion.adapter_a import DOMAIN_HEADERS, HEADER_TO_KEY

logger = logging.getLogger(__name__)

# Treatment-worded display headings written into the generated document, in order.
TREATMENT_HEADING_BY_KEY = {
    'רקע': 'רקע:',
    'מהלך_האבחון': 'מהלך הטיפול:',
    'ממצאי_האבחון': 'ממצאי הטיפול:',
    'סיכום_והמלצות': 'סיכום והמלצות:',
}

_ABSTRACT_KEY = 'תמצית_אבחון'          # section removed from the generated document
_SUBJECT_PREFIX = 'הנדון'
_DATE_PREFIX = 'תאריך'
_SIGNATURE_MARKERS = ('בברכה', 'מ.ר.')


def build_summary_docx(
    skeleton_path: str,
    sections: dict[str, str],
    domains: dict[str, str],
    header_date: str,
    header_fields: Optional[dict[str, str]] = None,
) -> Document:
    """Return a Document built from *skeleton_path* with the generated content injected.

    sections / domains hold the re-identified generated texts. header_fields
    (name / national_id / date_of_birth) is passed only for the clean bundled
    template, whose header holds empty placeholders — a copied patient document
    already carries real header values.
    """
    doc = Document(skeleton_path)

    _update_preamble(doc, header_date, header_fields)

    paragraphs = list(doc.paragraphs)
    headings = _locate_headings(paragraphs)

    for pos, (idx, key) in enumerate(headings):
        end = _body_end(paragraphs, headings, pos)
        body_range = paragraphs[idx + 1:end]

        if key == _ABSTRACT_KEY:
            for p in body_range:
                _delete_paragraph(p)
            _delete_paragraph(paragraphs[idx])
            continue

        if key not in TREATMENT_HEADING_BY_KEY:
            continue

        heading_p = paragraphs[idx]
        inline_body_rpr = _rewrite_heading(heading_p, TREATMENT_HEADING_BY_KEY[key])
        donor_ppr, donor_rpr = _pick_body_donor(heading_p, body_range, inline_body_rpr)

        for p in body_range:
            _delete_paragraph(p)

        if key == 'ממצאי_האבחון':
            _insert_findings(heading_p, sections.get(key, ''), domains, donor_ppr, donor_rpr)
        else:
            _insert_lines(heading_p, sections.get(key, ''), donor_ppr, donor_rpr)

    found_keys = {k for _, k in headings}
    missing = [k for k in TREATMENT_HEADING_BY_KEY if k not in found_keys]
    if missing:
        logger.warning("Skeleton %s is missing section headings: %s", skeleton_path, missing)

    return doc


# ---------------------------------------------------------------------------
# Preamble (subject line, date, template header placeholders)
# ---------------------------------------------------------------------------

def _update_preamble(doc, header_date: str, header_fields: Optional[dict]) -> None:
    for p in doc.paragraphs:
        if _heading_key(p) is not None:
            break
        line = _first_line(p)
        if line.startswith(_SUBJECT_PREFIX):
            _set_paragraph_text(p, line.replace('אבחון', 'טיפול'))
        elif line.startswith(_DATE_PREFIX):
            _set_paragraph_text(p, f'תאריך: {header_date}')
        elif header_fields is not None:
            _fill_id_fields(p, line, header_fields)


def _fill_id_fields(paragraph, line: str, fields: dict) -> None:
    """Fill the name / date-of-birth / national-id placeholders of the clean template.

    Handles both template layouts: three separate paragraphs (one field each), or a
    single paragraph holding all three fields separated by soft line breaks.
    """
    if line.startswith('שם'):
        if 'ת.ל' in paragraph.text or 'ת.ז' in paragraph.text:
            _fill_id_paragraph(paragraph, fields)
        else:
            _set_paragraph_text(paragraph, f"שם: {fields.get('name', '')}")
    elif line.startswith('ת.ל'):
        _set_paragraph_text(paragraph, f"ת.ל.: {fields.get('date_of_birth', '')}")
    elif line.startswith('ת.ז'):
        _set_paragraph_text(paragraph, f"ת.ז.: {fields.get('national_id', '')}")


def _fill_id_paragraph(paragraph, fields: dict) -> None:
    rpr = _run_rpr(paragraph.runs[0]) if paragraph.runs else None
    for run in list(paragraph.runs):
        run._r.getparent().remove(run._r)
    lines = [
        f"שם: {fields.get('name', '')}",
        f"ת.ל.: {fields.get('date_of_birth', '')}",
        f"ת.ז.: {fields.get('national_id', '')}",
    ]
    for i, line in enumerate(lines):
        run = _make_run(paragraph, line, rpr)
        if i < len(lines) - 1:
            run.add_break()


# ---------------------------------------------------------------------------
# Section location and boundaries
# ---------------------------------------------------------------------------

def _locate_headings(paragraphs) -> list[tuple[int, str]]:
    """Return [(paragraph_index, section_key)] for the first occurrence of each section."""
    found: list[tuple[int, str]] = []
    seen: set[str] = set()
    for idx, p in enumerate(paragraphs):
        key = _heading_key(p)
        if key is not None and key not in seen:
            found.append((idx, key))
            seen.add(key)
    return found


def _heading_key(paragraph) -> Optional[str]:
    line = _first_line(paragraph).rstrip(':').strip()
    if not line:
        return None
    for header_text, key in HEADER_TO_KEY.items():
        if line.startswith(header_text.rstrip(':')):
            return key
    return None


def _body_end(paragraphs, headings, pos: int) -> int:
    """Index one past the last body paragraph of the section at headings[pos]."""
    start = headings[pos][0] + 1
    next_heading = headings[pos + 1][0] if pos + 1 < len(headings) else len(paragraphs)
    for idx in range(start, next_heading):
        if _is_signature(paragraphs[idx]):
            return idx
    return next_heading


def _is_signature(paragraph) -> bool:
    return any(marker in paragraph.text for marker in _SIGNATURE_MARKERS)


def _first_line(paragraph) -> str:
    return paragraph.text.split('\n', 1)[0].strip()


# ---------------------------------------------------------------------------
# Content injection
# ---------------------------------------------------------------------------

def _rewrite_heading(paragraph, heading_text: str):
    """Reduce the paragraph to its (re-worded) heading only, preserving formatting.

    Template-style skeletons keep heading and body in one paragraph separated by
    soft line breaks; the rPr of the first inline body run is returned so injected
    body paragraphs can clone it.
    """
    runs = list(paragraph.runs)
    heading_rpr = _run_rpr(runs[0]) if runs else None

    inline_body_rpr = None
    break_pos = paragraph.text.find('\n')
    prefix = ''
    for run in runs:
        if break_pos != -1 and len(prefix) > break_pos and run.text.strip():
            inline_body_rpr = _run_rpr(run)
            break
        prefix += run.text

    for run in runs:
        run._r.getparent().remove(run._r)
    _make_run(paragraph, heading_text, heading_rpr)
    return inline_body_rpr


def _pick_body_donor(heading_p, body_range, inline_body_rpr):
    """Choose the pPr / rPr the injected body paragraphs will clone."""
    for p in body_range:
        if p.runs and p.text.strip():
            return _para_ppr(p), _run_rpr(p.runs[0])
    rpr = inline_body_rpr
    if rpr is None:
        rpr = _strip_bold(_run_rpr(heading_p.runs[0])) if heading_p.runs else None
    return _para_ppr(heading_p), rpr


def _insert_lines(ref_p, text: str, donor_ppr, donor_rpr, bold: Optional[bool] = None):
    """Insert one paragraph per non-empty line of *text* after *ref_p*; return the last."""
    last = ref_p
    for line in text.split('\n'):
        line = line.strip()
        if not line:
            continue
        last = _insert_paragraph_after(last, donor_ppr)
        _make_run(last, line, donor_rpr, bold=bold)
    return last


def _insert_findings(heading_p, intro_text: str, domains: dict, donor_ppr, donor_rpr) -> None:
    """Insert the findings intro, then each domain as a bold sub-heading + body."""
    last = _insert_lines(heading_p, intro_text, donor_ppr, donor_rpr)
    ordered = [d for d in DOMAIN_HEADERS if d in domains]
    ordered += [d for d in domains if d not in DOMAIN_HEADERS]
    for domain in ordered:
        text = (domains.get(domain) or '').strip()
        if not text:
            continue
        last = _insert_paragraph_after(last, donor_ppr)
        _make_run(last, f'{domain} – ', donor_rpr, bold=True)
        last = _insert_lines(last, text, donor_ppr, donor_rpr)


# ---------------------------------------------------------------------------
# python-docx low-level helpers
# ---------------------------------------------------------------------------

def _delete_paragraph(paragraph) -> None:
    paragraph._p.getparent().remove(paragraph._p)


def _set_paragraph_text(paragraph, text: str) -> None:
    rpr = _run_rpr(paragraph.runs[0]) if paragraph.runs else None
    for run in list(paragraph.runs):
        run._r.getparent().remove(run._r)
    _make_run(paragraph, text, rpr)


def _run_rpr(run):
    rpr = run._r.find(qn('w:rPr'))
    return copy.deepcopy(rpr) if rpr is not None else None


def _para_ppr(paragraph):
    ppr = paragraph._p.pPr
    return copy.deepcopy(ppr) if ppr is not None else None


def _strip_bold(rpr):
    if rpr is None:
        return None
    for tag in ('w:b', 'w:bCs'):
        el = rpr.find(qn(tag))
        if el is not None:
            rpr.remove(el)
    return rpr


def _make_run(paragraph, text: str, rpr, bold: Optional[bool] = None):
    run = paragraph.add_run(text)
    if rpr is not None:
        existing = run._r.find(qn('w:rPr'))
        if existing is not None:
            run._r.remove(existing)
        run._r.insert(0, copy.deepcopy(rpr))
    if bold is not None:
        run.bold = bold
    return run


def _insert_paragraph_after(ref_p: Paragraph, ppr) -> Paragraph:
    new_el = OxmlElement('w:p')
    if ppr is not None:
        new_el.append(copy.deepcopy(ppr))
    ref_p._p.addnext(new_el)
    return Paragraph(new_el, ref_p._parent)
