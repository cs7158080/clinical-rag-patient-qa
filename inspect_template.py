from docx import Document

template_path = r"D:\Users\Agents course\Lesson2 - RAG + Embedding\tamplates\טמפליט אבחון בנים.docx"
doc = Document(template_path)

print("=== TEMPLATE STRUCTURE ===\n")
line_num = 0
for para in doc.paragraphs:
    text = para.text.strip()
    if text:
        print(f"{line_num}: {text}")
        line_num += 1

print("\n=== TABLES ===")
for t_idx, table in enumerate(doc.tables):
    print(f"\nTable {t_idx}: {len(table.rows)} rows x {len(table.columns)} cols")
    for r_idx, row in enumerate(table.rows[:5]):
        cells = [cell.text.strip()[:40] for cell in row.cells]
        print(f"  Row {r_idx}: {cells}")
