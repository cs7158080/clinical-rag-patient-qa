"""
audit_egress.py — Egress audit: prove no PHI crosses the network boundary.

Wraps every outbound client (Anthropic sync SDK, LlamaIndex Anthropic,
Cohere embed, Pinecone) so each payload is captured and scanned BEFORE send,
then drives every flow: full ingestion, one query per routing strategy, and
summary generation.

Modes
-----
Default        : mock — intercepted clients return fake responses (zero cost,
                 no network) + synthetic patient in a temp dir. Real data/ is
                 never touched.
--live         : forward captured calls to the real providers (costs money;
                 synthetic vectors upserted to the real Pinecone index are
                 deleted at the end of the run).
--real-corpus  : READ-ONLY — run only the query flows against the real DB;
                 forbidden list built from the real reid_map values + patient
                 folder names. No ingestion, no generation, no writes.

Requires the NER ONNX model (run setup.bat first) — the S5-2 proof depends
on NER catching a third-party name; without the model the audit refuses to
run (except --real-corpus, which performs no ingestion).

Honest scope note: this audit proves that no forbidden value crossed the SDK
boundary IN THE EXERCISED FLOWS — which is why it deliberately exercises all
five routing strategies and every flow. It cannot prove a negative for code
paths that never ran.

Usage:
    uv run python scripts/audit_egress.py [--live] [--real-corpus]

Exit codes: 0 = PASS, 1 = findings detected, 2 = setup error.
"""

import argparse
import asyncio
import hashlib
import json
import logging
import os
import re
import shutil
import sys
import tempfile
from dataclasses import replace
from types import SimpleNamespace

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

# Windows console defaults to cp1255 — Hebrew findings must still print
if hasattr(sys.stdout, "reconfigure"):
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")

# ---------------------------------------------------------------------------
# Synthetic identity — entirely fictitious.
# NOTE: the name-variant validator requires the LAST whitespace token to match
# across folder name, filename, and header — keep one consistent order.
# ---------------------------------------------------------------------------

SYN_FOLDER = "דנה ישראלי"
SYN_THIRD_PARTY = "רות אברמוביץ"
SYN_PHONES = ["0521234567", "052-1234567", "052 123 4567", "+972-52-1234567", "03-1234567"]
SYN_NATIONAL_IDS = ["123456789", "12345678-9", "23456781"]
SYN_DOB = "01/05/2019"
SYN_EMAIL = "dana.parent@example.com"

MOCK_QA_ANSWER = "על פי המסמכים חלה התקדמות עקבית בתחומים שנבדקו."
MOCK_GENERATION_JSON = json.dumps({
    "רקע": "רקע מעודכן.",
    "מהלך_האבחון": "מהלך הטיפול בתקופה.",
    "ממצאי_האבחון": "להלן הממצאים לפי תחומים.",
    "סיכום_והמלצות": "מומלץ המשך טיפול.",
    "ממצאי_תחומים": {"זיכרון שמיעתי": "שיפור עקבי.", "הבנת שפה": "ללא שינוי."},
}, ensure_ascii=False)

# (question, strategy, canned extraction params) — one per routing strategy
QUERY_MATRIX = [
    ("תסכמי את האבחון",
     "family_a_sqlite",
     {"intent": "summarize", "template_type": "diagnosis", "topic": None,
      "date_from": None, "date_to": None, "session_count": None}),
    ("תסכמי את הטיפולים",
     "treatment_sessions_sqlite",
     {"intent": "summarize", "template_type": None, "topic": None,
      "date_from": None, "date_to": None, "session_count": None}),
    ("מה מצב הזיכרון השמיעתי?",
     "domain_sqlite",
     {"intent": "check_domain", "template_type": None, "topic": "זיכרון שמיעתי",
      "date_from": None, "date_to": None, "session_count": None}),
    ("האם עבדנו על קשב והתמדה?",
     "pinecone",
     {"intent": "find_specific", "template_type": None, "topic": "קשב והתמדה",
      "date_from": None, "date_to": None, "session_count": None}),
    ("האם חל שיפור מאז תחילת החודש?",
     "compare_sqlite",
     {"intent": "compare_progress", "template_type": None, "topic": None,
      "date_from": "2026-03-08", "date_to": None, "session_count": None}),
]

_HEX64 = re.compile(r"[0-9a-f]{64}")
_TOKEN_SCRUB = re.compile(r"(?:(?:PERSON|INST|ID|PHONE|EMAIL)_)?[0-9a-f]{64}")


# ---------------------------------------------------------------------------
# Auditor — scan every captured payload
# ---------------------------------------------------------------------------

class EgressAuditor:
    def __init__(self, log_path: str):
        self.flow = "setup"
        self.failures: "list[str]" = []
        self.payload_count = 0
        self._name_patterns: "list[tuple[str, re.Pattern]]" = []
        self._digit_values: "list[tuple[str, str]]" = []
        self._plain_values: "list[str]" = []
        self._forbidden_hashes: "set[str]" = set()
        self._log = logging.getLogger("egress_audit")
        handler = logging.FileHandler(log_path, mode="w", encoding="utf-8")
        handler.setFormatter(logging.Formatter("%(asctime)s %(message)s"))
        self._log.addHandler(handler)
        self._log.setLevel(logging.INFO)
        self._log.propagate = False

    # -- forbidden list construction ------------------------------------
    def add_name(self, name: str) -> None:
        """Hebrew name: word-boundary match with one optional attached prefix
        (catches ודנה/לרות without false-firing on שירות)."""
        for part in {name} | set(name.split()):
            if len(part) < 2:
                continue
            pattern = re.compile(
                rf"(?<![\w֐-׿])[ולבמהשכ]?{re.escape(part)}(?![\w֐-׿])"
            )
            self._name_patterns.append((part, pattern))
            self._forbidden_hashes.add(hashlib.sha256(part.encode()).hexdigest())

    def add_digits(self, label: str, value: str) -> None:
        self._digit_values.append((label, re.sub(r"[\s\-./]", "", value)))

    def add_plain(self, value: str) -> None:
        if value and len(value) >= 4:
            self._plain_values.append(value)

    # -- scanning ---------------------------------------------------------
    def scan(self, destination: str, payload) -> None:
        from app.deidentification.patterns import (
            EMAIL_REGEX, NATIONAL_ID_REGEX, PHONE_REGEX,
        )

        self.payload_count += 1
        findings: "list[str]" = []
        for raw in _iter_strings(payload):
            # Layer 0 — S5-4: no captured 64-hex may equal sha256(known name)
            for hex_str in _HEX64.findall(raw):
                if hex_str in self._forbidden_hashes:
                    findings.append(f"sha256-derivable key detected: {hex_str[:12]}…")
            # Tokens/keys are trusted 64-hex — mask them so their random
            # digit runs can't false-trip the digit/regex layers below.
            text = _TOKEN_SCRUB.sub(" ", raw)
            # Layer 1 — names (boundary + optional prefix)
            for name, pattern in self._name_patterns:
                if pattern.search(text):
                    findings.append(f"name '{name}' found")
            # Layer 2 — digit values, separator-normalized on both sides
            norm = re.sub(r"[\s\-./]", "", text)
            for label, digits in self._digit_values:
                if digits and digits in norm:
                    findings.append(f"digits of {label} found")
            # Layer 3 — plain substrings (email, raw map values)
            for value in self._plain_values:
                if value in text:
                    findings.append(f"forbidden value found: {value[:6]}…")
            # Layer 4 — S5-1 structural regexes as a safety net
            for label, rx in (("phone", PHONE_REGEX), ("national_id", NATIONAL_ID_REGEX),
                              ("email", EMAIL_REGEX)):
                match = re.search(rx, text)
                if match:
                    findings.append(f"regex '{label}' matched: {match.group(0)}")
        verdict = "FAIL" if findings else "ok"
        self._log.info("flow=%s dest=%s verdict=%s details=%s",
                       self.flow, destination, verdict, findings or "-")
        for f in findings:
            self.failures.append(f"[{self.flow} → {destination}] {f}")

    def report(self) -> int:
        print(f"\n=== EGRESS AUDIT: {self.payload_count} payloads scanned ===")
        if not self.failures:
            print("PASS — no forbidden value crossed the network boundary "
                  "in any exercised flow.")
            return 0
        print(f"FAIL — {len(self.failures)} finding(s):")
        for f in self.failures:
            print(f"  {f}")
        return 1


def _iter_strings(obj):
    """Yield every string in a nested payload (dict keys+values, lists).
    Non-string scalars (floats/ints — embedding values, date nums) are skipped."""
    if isinstance(obj, str):
        yield obj
    elif isinstance(obj, dict):
        for k, v in obj.items():
            yield from _iter_strings(k)
            yield from _iter_strings(v)
    elif isinstance(obj, (list, tuple, set)):
        for item in obj:
            yield from _iter_strings(item)


# ---------------------------------------------------------------------------
# Interception — patch each client at its egress boundary
# ---------------------------------------------------------------------------

def install_interceptors(auditor: EgressAuditor, live: bool, dim: int):
    from anthropic.resources.messages import Messages
    from llama_index.llms.anthropic import Anthropic as LlamaAnthropic
    from llama_index.core.llms import CompletionResponse
    import cohere

    extraction_table = {q: params for q, _, params in QUERY_MATRIX}

    orig_create = Messages.create

    def patched_create(self, **kwargs):
        # "model" is a constant provider id (e.g. ...-20251001) whose 8-digit
        # date false-trips the national-id regex — it is not user data.
        auditor.scan("anthropic.messages.create",
                     {k: v for k, v in kwargs.items() if k != "model"})
        if live:
            return orig_create(self, **kwargs)
        question = ""
        msgs = kwargs.get("messages") or []
        if msgs and isinstance(msgs[0].get("content"), str):
            question = msgs[0]["content"]
        params = extraction_table.get(question)
        text = json.dumps(params, ensure_ascii=False) if params else MOCK_QA_ANSWER
        return SimpleNamespace(content=[SimpleNamespace(text=text)])

    Messages.create = patched_create

    orig_acomplete = LlamaAnthropic.acomplete

    async def patched_acomplete(self, prompt, **kwargs):
        auditor.scan("llamaindex.anthropic.acomplete", {"prompt": prompt})
        if live:
            return await orig_acomplete(self, prompt, **kwargs)
        text = MOCK_GENERATION_JSON if "ממצאי_תחומים" in prompt else MOCK_QA_ANSWER
        return CompletionResponse(text=text)

    LlamaAnthropic.acomplete = patched_acomplete

    orig_embed = cohere.ClientV2.embed

    def patched_embed(self, *, texts, model, input_type, embedding_types=None, **kwargs):
        auditor.scan("cohere.embed", {"texts": list(texts)})
        if live:
            return orig_embed(self, texts=texts, model=model, input_type=input_type,
                              embedding_types=embedding_types, **kwargs)
        return SimpleNamespace(
            embeddings=SimpleNamespace(float_=[[0.01] * dim for _ in texts])
        )

    cohere.ClientV2.embed = patched_embed


class FakeIndex:
    """Mock-mode Pinecone index: captures payloads, serves stored vectors back."""

    def __init__(self, auditor: EgressAuditor):
        self._auditor = auditor
        self._vectors: "dict[str, dict]" = {}

    def upsert(self, vectors):
        self._auditor.scan("pinecone.upsert", vectors)
        for v in vectors:
            self._vectors[v["id"]] = v

    def query(self, vector=None, filter=None, top_k=10, include_metadata=True):
        self._auditor.scan("pinecone.query", {"filter": filter, "top_k": top_k})
        pid = (filter or {}).get("patient_id", {}).get("$eq")
        matches = [
            SimpleNamespace(id=vid, score=0.9, metadata=v["metadata"])
            for vid, v in self._vectors.items()
            if v["metadata"].get("patient_id") == pid
        ][:top_k]
        return SimpleNamespace(matches=matches)

    def delete(self, ids=None, delete_all=False):
        self._auditor.scan("pinecone.delete", {"ids": ids or []})
        for vid in ids or []:
            self._vectors.pop(vid, None)


class CaptureIndex:
    """Live-mode proxy: captures payloads, forwards to the real index, and
    remembers upserted ids so the run can clean them up afterwards."""

    def __init__(self, auditor: EgressAuditor, real_index):
        self._auditor = auditor
        self._real = real_index
        self.upserted_ids: "list[str]" = []

    def upsert(self, vectors):
        self._auditor.scan("pinecone.upsert", vectors)
        self.upserted_ids.extend(v["id"] for v in vectors)
        return self._real.upsert(vectors=vectors)

    def query(self, **kwargs):
        self._auditor.scan("pinecone.query",
                           {k: v for k, v in kwargs.items() if k != "vector"})
        return self._real.query(**kwargs)

    def delete(self, ids=None, delete_all=False):
        self._auditor.scan("pinecone.delete", {"ids": ids or []})
        return self._real.delete(ids=ids)

    def cleanup(self):
        if self.upserted_ids:
            self._real.delete(ids=self.upserted_ids)
            print(f"Cleaned {len(self.upserted_ids)} synthetic vectors from the real index.")


# ---------------------------------------------------------------------------
# Synthetic corpus
# ---------------------------------------------------------------------------

def build_synthetic_corpus(patients_root: str) -> None:
    from docx import Document

    folder = os.path.join(patients_root, SYN_FOLDER)
    os.makedirs(folder)

    doc = Document()
    for line in (
        f"שם: {SYN_FOLDER}", "תאריך: 01/02/2026", f"ת.ל.: {SYN_DOB}",
        f"ת.ז.: {SYN_NATIONAL_IDS[0]}", "קופת חולים: מכבי",
        "רקע",
        f"דנה הופנתה לאבחון על ידי ההורים. טלפון ליצירת קשר: {SYN_PHONES[1]} או {SYN_PHONES[2]}.",
        f"בשעת הצורך ניתן לפנות למספר {SYN_PHONES[3]}, לקו הקווי {SYN_PHONES[4]} "
        f"או בדוא\"ל {SYN_EMAIL}.",
        f"מספר תעודת הזהות של דנה: {SYN_NATIONAL_IDS[1]}, ובמסמכים ישנים נרשם {SYN_NATIONAL_IDS[2]}.",
        # long filler — pushes the next line far past the old 512-token horizon (S5-2)
        "במהלך המפגשים נצפו שיתוף פעולה מלא והתקדמות עקבית בתחומי השפה, "
        "התקשורת והקשב, לצד הנאה ניכרת מהמשחק המשותף. " * 60,
        f"בשיחת הסיכום השתתפה גם הסייעת {SYN_THIRD_PARTY}, שמלווה את דנה בגן. "
        f"ודנה שמחה מאוד על נוכחותה.",
        "מהלך האבחון",
        "האבחון נערך בשני מפגשים וכלל תצפית, משחק מובנה ומטלות שפה.",
        "ממצאי האבחון",
        "להלן הממצאים לפי תחומים.",
        "זיכרון שמיעתי",
        "חזרה על רצפים של שלושה פריטים ללא קושי.",
        "הבנת שפה",
        "הבנת הוראות מורכבות תואמת גיל.",
        "סיכום והמלצות",
        "מומלץ טיפול שבועי לחיזוק הזיכרון השמיעתי.",
        "בברכה,",
        "קלינאית תקשורת, מ.ר. 54321",
    ):
        doc.add_paragraph(line)
    doc.save(os.path.join(folder, f"סיכום אבחון {SYN_FOLDER}.docx"))

    doc = Document()
    for line in (
        f"שם: {SYN_FOLDER}", f"ת.ל.: {SYN_DOB}",
        "תאריך התחלת הטיפול: 01/03/2026",
        "מטרת על: שיפור הזיכרון השמיעתי והבעת השפה",
    ):
        doc.add_paragraph(line)
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "01/03/2026"
    table.rows[0].cells[1].text = "חיזוק זיכרון שמיעתי באמצעות רצפי מספרים"
    table.rows[1].cells[0].text = "15/03/2026"
    table.rows[1].cells[1].text = "הרחבת אוצר מילים בקטגוריות משחק"
    for line in (
        "01/03/2026",
        f"עבדנו עם דנה על רצפי מספרים; ודנה חזרה על ארבעה פריטים. "
        f"ההורה עדכן מספר טלפון חדש: {SYN_PHONES[0]}.",
        "15/03/2026",
        f"תרגול קטגוריות בהשתתפות הסייעת {SYN_THIRD_PARTY}. נצפה שיפור עקבי אצל דנה.",
    ):
        doc.add_paragraph(line)
    doc.save(os.path.join(folder, f"תוכנית טיפול {SYN_FOLDER}.docx"))


def build_synthetic_forbidden_list(auditor: EgressAuditor) -> None:
    for name in (SYN_FOLDER, SYN_THIRD_PARTY):
        auditor.add_name(name)
    for i, phone in enumerate(SYN_PHONES):
        auditor.add_digits(f"phone[{i}]", phone)
    for i, nid in enumerate(SYN_NATIONAL_IDS):
        auditor.add_digits(f"national_id[{i}]", nid)
    auditor.add_digits("date_of_birth", SYN_DOB)
    auditor.add_plain(SYN_EMAIL)


# ---------------------------------------------------------------------------
# Flows
# ---------------------------------------------------------------------------

async def run_flows_synthetic(auditor, config, db_path, reid_map_path, index):
    from app.ingestion.pipeline import run_ingestion
    from app.generation.qa import run_query
    from app.generation.summary_generator import run_generate_summary
    from app.deidentification.reid_map import load as load_reid_map, patient_id_from_folder

    auditor.flow = "ingestion"
    reid_map = load_reid_map(reid_map_path)
    results = await run_ingestion(
        patients_roots=config.patients_roots, config=config,
        reid_map=reid_map, db_path=db_path, pinecone_index=index,
    )
    for path, result in results:
        print(f"  ingest {result:20s} {os.path.basename(path)}")
        if result != "ok":
            auditor.failures.append(f"[ingestion] file not ingested cleanly: "
                                    f"{os.path.basename(path)} → {result}")

    reid_map = load_reid_map(reid_map_path)
    patient_id = patient_id_from_folder(reid_map, SYN_FOLDER)
    if patient_id is None:
        auditor.failures.append("[setup] synthetic patient missing from reid_map after ingest")
        return

    for question, strategy, _ in QUERY_MATRIX:
        auditor.flow = f"query:{strategy}"
        answer = await run_query(question, patient_id, config, db_path, index, reid_map_path)
        print(f"  query  {strategy:28s} answered ({len(answer)} chars)")

    auditor.flow = "generation"
    result = await run_generate_summary(
        patient_id, "2026-03-01", "2026-03-15", config, db_path, reid_map_path,
    )
    print(f"  generation → {result[:80]}")


async def run_flows_real_corpus(auditor, config, db_path, reid_map_path, index):
    """READ-ONLY: query flows only, against the real DB."""
    from app.generation.qa import run_query
    from app.storage import db

    patients = db.get_patient_list(db_path)
    if not patients:
        auditor.failures.append("[setup] real DB has no patients")
        return
    patient_id = patients[0]["patient_id"]
    dates = db.get_treatment_session_dates(db_path, patient_id)
    mid_date = dates[len(dates) // 2] if dates else None

    for question, strategy, params in QUERY_MATRIX:
        if strategy == "compare_sqlite" and mid_date:
            params["date_from"] = mid_date
        auditor.flow = f"query:{strategy}"
        answer = await run_query(question, patient_id, config, db_path, index, reid_map_path)
        print(f"  query  {strategy:28s} answered ({len(answer)} chars)")


def build_real_forbidden_list(auditor, config, reid_map_path) -> None:
    from app.deidentification.reid_map import load as load_reid_map

    for value in load_reid_map(reid_map_path).values():
        if not value:
            continue
        if re.fullmatch(r"[\d\s\-./+]+", value):
            auditor.add_digits("reid_map digit value", value)
        elif re.search(r"[֐-׿]", value):
            auditor.add_name(value)
        else:
            auditor.add_plain(value)
    for patients_root in config.patients_roots:
        if not os.path.isdir(patients_root):
            continue
        for entry in os.scandir(patients_root):
            if entry.is_dir():
                auditor.add_name(entry.name)


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main() -> int:
    parser = argparse.ArgumentParser(description="Egress audit (see module docstring)")
    parser.add_argument("--live", action="store_true",
                        help="forward captured calls to the real providers")
    parser.add_argument("--real-corpus", action="store_true",
                        help="read-only query audit against the real DB")
    args = parser.parse_args()

    if not args.live:
        for key in ("ANTHROPIC_API_KEY", "COHERE_API_KEY", "PINECONE_API_KEY"):
            os.environ.setdefault(key, "mock-key-never-sent")

    from app.config import get_config
    config = get_config()

    os.makedirs(config.logs_dir, exist_ok=True)
    log_path = os.path.join(config.logs_dir, "egress_audit.log")
    auditor = EgressAuditor(log_path)
    # logging.basicConfig(level=logging.WARNING)

    tmp_root = None
    capture_index = None
    try:
        if args.real_corpus:
            db_path = os.path.join(config.data_dir, "clinical_rag.db")
            reid_map_path = os.path.join(config.data_dir, "reid_map.json")
            if not os.path.isfile(db_path):
                print(f"Real DB not found: {db_path}")
                return 2
            build_real_forbidden_list(auditor, config, reid_map_path)
            run_config = config
        else:
            # NER is mandatory for the synthetic run (S5-2 proof needs it)
            from app.deidentification.ner import load_ner_model
            try:
                load_ner_model(config.models_dir)
            except FileNotFoundError as exc:
                print(f"NER model required for the audit but not available:\n{exc}")
                return 2

            tmp_root = tempfile.mkdtemp(prefix="egress_audit_")
            patients_root = os.path.join(tmp_root, "patients")
            data_dir = os.path.join(tmp_root, "data")
            os.makedirs(patients_root)
            os.makedirs(data_dir)
            build_synthetic_corpus(patients_root)
            build_synthetic_forbidden_list(auditor)

            run_config = replace(config, patients_roots=[patients_root], data_dir=data_dir)
            db_path = os.path.join(data_dir, "clinical_rag.db")
            reid_map_path = os.path.join(data_dir, "reid_map.json")
            from app.storage.db import init_db
            init_db(db_path)

        install_interceptors(auditor, live=args.live, dim=config.pinecone.dimension)

        if args.live:
            from app.storage.pinecone_client import init_pinecone
            real_index = init_pinecone(config.pinecone_api_key, config.pinecone.index_name)
            capture_index = CaptureIndex(auditor, real_index)
            index = capture_index
        else:
            index = FakeIndex(auditor)

        if args.real_corpus:
            asyncio.run(run_flows_real_corpus(auditor, run_config, db_path,
                                              reid_map_path, index))
        else:
            asyncio.run(run_flows_synthetic(auditor, run_config, db_path,
                                            reid_map_path, index))
    finally:
        if capture_index is not None:
            capture_index.cleanup()

    exit_code = auditor.report()
    print(f"Full log: {log_path}")
    if tmp_root:
        if exit_code == 0:
            shutil.rmtree(tmp_root, ignore_errors=True)
        else:
            print(f"Synthetic corpus kept for inspection: {tmp_root}")
    return exit_code


if __name__ == "__main__":
    sys.exit(main())
