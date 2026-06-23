"""
Test script for src/chunker.py

Creates a temporary .docx with Hebrew content matching the 'diagnosis' template,
runs Chunker.chunk_file(), prints results, and asserts basic correctness.
"""

import os
import sys
import tempfile

# Make sure we can import from src/ even when run from the project root
sys.path.insert(0, os.path.dirname(__file__))

from docx import Document
from src.chunker import Chunker, TemplateRegistry

# ---------------------------------------------------------------------------
# Build a minimal in-memory .docx that matches the diagnosis template
# ---------------------------------------------------------------------------

def build_test_docx(path: str) -> None:
    doc = Document()

    # Preamble metadata fields (not section headers)
    doc.add_paragraph("תאריך: 01/06/2025")
    doc.add_paragraph("ת.ל.: 15/03/2000")

    # Section 1
    doc.add_paragraph("רקע:")
    doc.add_paragraph("הילד נולד בשנת 2000 ומתגורר עם משפחתו בתל אביב.")
    doc.add_paragraph("הפנייה לאבחון בוצעה על ידי המחנכת בכיתה ג.")
    doc.add_paragraph("ההורים מדווחים על קשיים בריכוז ובקשב.")

    # Section 2
    doc.add_paragraph("מהלך האבחון:")
    doc.add_paragraph("האבחון נמשך שלושה מפגשים של שעה כל אחד.")
    doc.add_paragraph("בוצעו מבחני אינטליגנציה, קריאה וכתיבה.")
    doc.add_paragraph("שיתוף הפעולה של הנבחן היה טוב.")

    # Section 3
    doc.add_paragraph("ממצאי האבחון:")
    doc.add_paragraph("רמת האינטליגנציה הכללית נמצאת בטווח הממוצע.")
    doc.add_paragraph("זוהו קשיים בזיכרון פעיל ובעיבוד שמיעתי.")
    doc.add_paragraph("הקריאה שוטפת אך הכתיבה לוקה בהשמטות.")

    # Section 4
    doc.add_paragraph("סיכום והמלצות:")
    doc.add_paragraph("מומלץ על תגבור לימודי בקריאה וכתיבה.")
    doc.add_paragraph("יש לשקול הפניה לריפוי בעיסוק לבדיקת כישורים מוטוריים.")
    doc.add_paragraph("מומלץ להתאים תנאי בחינה מיוחדים.")

    doc.save(path)


# ---------------------------------------------------------------------------
# Main test
# ---------------------------------------------------------------------------

def main():
    # Determine the project root (where templates_config.yaml lives)
    project_root = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    config_path = os.path.join(project_root, "templates_config.yaml")

    # Create temp .docx with an ASCII filename to avoid Windows encoding issues
    tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False, prefix="test_client_")
    tmp_path = tmp.name
    tmp.close()  # close so python-docx can write to it on Windows

    try:
        build_test_docx(tmp_path)
        print(f"Temp file: {tmp_path}\n")

        registry = TemplateRegistry(config_path=config_path)
        chunker = Chunker(registry=registry)
        chunks = chunker.chunk_file(tmp_path)

        print(f"Total chunks: {len(chunks)}\n")
        print("-" * 60)
        for i, chunk in enumerate(chunks):
            section = chunk.metadata["section_title"]
            preview = chunk.text[:80].replace("\n", " ")
            print(f"[{i}] section_title={section!r}")
            print(f"     chunk_index={chunk.metadata['chunk_index']}  "
                  f"char_count={chunk.metadata['char_count']}  "
                  f"is_free_text={chunk.metadata['is_free_text']}")
            print(f"     text preview: {preview!r}")
            print()

        # --- Assertions ---
        assert len(chunks) > 0, "Expected at least one chunk but got 0"

        section_titles = [c.metadata["section_title"] for c in chunks]
        assert "רקע:" in section_titles, "Expected section 'רקע:' to appear in chunks"
        assert "מהלך האבחון:" in section_titles, "Expected section 'מהלך האבחון:' to appear in chunks"
        assert "ממצאי האבחון:" in section_titles, "Expected section 'ממצאי האבחון:' to appear in chunks"
        assert "סיכום והמלצות:" in section_titles, "Expected section 'סיכום והמלצות:' to appear in chunks"

        # template metadata
        assert chunks[0].metadata["template_type"] == "diagnosis"
        assert chunks[0].metadata["template_name"] == "אבחון"

        # date fields extracted from preamble
        dates = [c.metadata["document_date"] for c in chunks]
        assert all(d == "01/06/2025" for d in dates), f"document_date mismatch: {dates}"
        birth_dates = [c.metadata["birth_date"] for c in chunks]
        assert all(b == "15/03/2000" for b in birth_dates), f"birth_date mismatch: {birth_dates}"

        # client_name comes from filename stem
        stem = os.path.splitext(os.path.basename(tmp_path))[0]
        assert all(c.metadata["client_name"] == stem for c in chunks)

        # --- BUG-001: node ID must incorporate content so updates are detected ---
        from src.chunker import Chunk
        meta = {"source": "/fake/path.docx", "chunk_index": 0}
        node_orig = Chunk(text="Original content", metadata=meta).to_llama_node()
        node_updated = Chunk(text="Updated content", metadata=meta).to_llama_node()
        node_same = Chunk(text="Original content", metadata=meta).to_llama_node()
        assert node_orig.id_ != node_updated.id_, "Same path+index but different text must yield different node IDs"
        assert node_orig.id_ == node_same.id_, "Identical chunks must yield identical node IDs (deterministic)"

        print("=" * 60)
        print("ALL ASSERTIONS PASSED")

    finally:
        os.unlink(tmp_path)


if __name__ == "__main__":
    main()
