"""
create_synthetic_docs.py — Generate synthetic .docx test fixtures.

Run once to produce test fixtures in tests/synthetic_data/.
All PII is entirely fictitious. No real patient data is used.

Files created:
    diagnosis_test.docx          — Family A, diagnosis template
    clinic_visit_summary_test.docx — Family A, clinic_visit_summary template
    treatment_plan_test.docx     — Family B, treatment_plan template
                                   (3 goals rows + 5 session blocks)

Patient:  יוסי כהן (first name last name ordering inside doc)
Folder:   כהן יוסי (last name first — folder naming convention)
"""

import os
from docx import Document
from docx.shared import Pt


# ---------------------------------------------------------------------------
# Constants — all fictitious data
# ---------------------------------------------------------------------------

PATIENT_NAME_IN_DOC = "יוסי כהן"
PATIENT_FOLDER_NAME = "כהן יוסי"  # last name first, used for folder

NATIONAL_ID = "123456789"
INSTITUTION = "בית ספר אלון"
PHONE = "0521234567"
EMAIL = "test@test.com"
DATE_OF_BIRTH = "01/05/2018"
HMO = "מכבי"

OUTPUT_DIR = os.path.dirname(os.path.abspath(__file__))


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _heading(doc: Document, text: str) -> None:
    """Add a bold paragraph as a section heading."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)


def _field_row(doc: Document, label: str, value: str) -> None:
    """Add a 'label: value' paragraph."""
    p = doc.add_paragraph()
    run_label = p.add_run(f"{label}: ")
    run_label.bold = True
    p.add_run(value)


# ---------------------------------------------------------------------------
# diagnosis_test.docx
# ---------------------------------------------------------------------------

def create_diagnosis_doc(output_path: str) -> None:
    """Create a synthetic diagnosis .docx with all PII types present."""
    doc = Document()

    # --- Header fields ---
    doc.add_heading("אבחון בנים — דוח אבחון", level=1)
    _field_row(doc, "שם", PATIENT_NAME_IN_DOC)
    _field_row(doc, "ת.ז.", NATIONAL_ID)
    _field_row(doc, "ת.ל.", DATE_OF_BIRTH)
    _field_row(doc, "קופת חולים", HMO)
    _field_row(doc, "תאריך", "15/01/2025")
    _field_row(doc, "טלפון", PHONE)
    _field_row(doc, "דוא\"ל", EMAIL)
    _field_row(doc, "מוסד חינוכי", INSTITUTION)
    doc.add_paragraph()

    # --- Section: רקע ---
    _heading(doc, "רקע")
    doc.add_paragraph(
        f"{PATIENT_NAME_IN_DOC} הוא ילד בן שבע. הוא לומד ב{INSTITUTION}. "
        f"ת.ז. {NATIONAL_ID}. ניתן ליצור קשר בטלפון {PHONE} "
        f"או במייל {EMAIL}."
    )

    # --- Section: מהלך האבחון ---
    _heading(doc, "מהלך האבחון")
    doc.add_paragraph(
        "האבחון בוצע בשני מפגשים. הילד שיתף פעולה בצורה טובה ופגין מוטיבציה גבוהה לאורך כל הפגישות."
    )

    # --- Section: ממצאי האבחון ---
    _heading(doc, "ממצאי האבחון")
    doc.add_paragraph(
        "פרגמטיקה ותקשורת: יכולות תקשורת בסיסיות תקינות. "
        "זיכרון שמיעתי: קשיים בזיכרון שמיעתי לטווח קצר. "
        "מודעות פונולוגית: מתחת לנורמה לגיל."
    )

    # --- Section: סיכום והמלצות ---
    _heading(doc, "סיכום והמלצות")
    doc.add_paragraph(
        "מומלץ על טיפול בקליניקה פעמיים בשבוע. ההורים יקבלו הנחיות לעבודה בבית."
    )

    # --- Section: תמצית אבחון ---
    _heading(doc, "תמצית אבחון")
    doc.add_paragraph(
        "ילד עם קשיי שפה ברמה קלה-בינונית. ניתן לצפות לשיפור משמעותי עם טיפול מתאים."
    )

    # --- Clinician signature block (should be stripped by adapter) ---
    doc.add_paragraph()
    doc.add_paragraph("___________________________")
    doc.add_paragraph("ד\"ר שרה לוי, קלינאית תקשורת")
    doc.add_paragraph("מ.ר. 12345")

    doc.save(output_path)
    print(f"Created: {output_path}")


# ---------------------------------------------------------------------------
# clinic_visit_summary_test.docx
# ---------------------------------------------------------------------------

def create_clinic_visit_summary_doc(output_path: str) -> None:
    """Create a synthetic clinic_visit_summary .docx with all PII types present."""
    doc = Document()

    doc.add_heading("סיכום ביקור 1 יוסי כהן", level=1)
    _field_row(doc, "שם", PATIENT_NAME_IN_DOC)
    _field_row(doc, "ת.ז.", NATIONAL_ID)
    _field_row(doc, "ת.ל.", DATE_OF_BIRTH)
    _field_row(doc, "קופת חולים", HMO)
    _field_row(doc, "תאריך", "10/03/2025")
    _field_row(doc, "טלפון", PHONE)
    _field_row(doc, "דוא\"ל", EMAIL)
    _field_row(doc, "מוסד חינוכי", INSTITUTION)
    doc.add_paragraph()

    _heading(doc, "רקע")
    doc.add_paragraph(
        f"ביקור ראשון עם {PATIENT_NAME_IN_DOC} לאחר תהליך האבחון. "
        f"הילד מגיע מ{INSTITUTION}."
    )

    _heading(doc, "מהלך האבחון")
    doc.add_paragraph(
        "בפגישה זו התמקדנו בהערכה מחדש של כישורי הדיבור. "
        "הילד שיתף פעולה היטב."
    )

    _heading(doc, "ממצאי האבחון")
    doc.add_paragraph(
        "שיפור ניכר בפרגמטיקה. זיכרון שמיעתי עדיין מצריך חיזוק."
    )

    _heading(doc, "סיכום והמלצות")
    doc.add_paragraph(
        "להמשיך טיפול שבועי. לעבוד על תרגילי זיכרון שמיעתי בבית."
    )

    _heading(doc, "תמצית אבחון")
    doc.add_paragraph(
        "התקדמות טובה. המשך טיפול מומלץ."
    )

    doc.add_paragraph()
    doc.add_paragraph("___________________________")
    doc.add_paragraph("ד\"ר שרה לוי, קלינאית תקשורת")
    doc.add_paragraph("מ.ר. 12345")

    doc.save(output_path)
    print(f"Created: {output_path}")


# ---------------------------------------------------------------------------
# treatment_plan_test.docx
# ---------------------------------------------------------------------------

def create_treatment_plan_doc(output_path: str) -> None:
    """Create a synthetic treatment_plan .docx.

    Layer 1: Goals table with 3 rows.
    Layer 2: 5 session blocks, each starting with a date line.
    """
    doc = Document()

    doc.add_heading("תוכנית טיפול הדרכתי יוסי כהן", level=1)
    _field_row(doc, "שם", PATIENT_NAME_IN_DOC)
    _field_row(doc, "ת.ל.", DATE_OF_BIRTH)
    _field_row(doc, "תאריך התחלת הטיפול", "01/01/2025")
    _field_row(doc, "מטרת על", "שיפור תקשורת ושפה")
    _field_row(doc, "הערות", "טיפול מוגבר בזמן הלימודים")
    doc.add_paragraph()

    # --- Layer 1: Goals table (3 rows) ---
    _heading(doc, "מטרות טיפוליות")
    table = doc.add_table(rows=4, cols=2)
    table.style = "Table Grid"

    # Header row
    header_cells = table.rows[0].cells
    header_cells[0].text = "תאריך"
    header_cells[1].text = "מטרות"

    # Row 1
    row1 = table.rows[1].cells
    row1[0].text = "01/01/2025"
    row1[1].text = (
        "1. שיפור זיכרון שמיעתי לטווח קצר\n"
        "2. עבודה על מודעות פונולוגית\n"
        "3. הרחבת אוצר מילים"
    )

    # Row 2
    row2 = table.rows[2].cells
    row2[0].text = "01/02/2025"
    row2[1].text = (
        "1. המשך עבודה על זיכרון שמיעתי\n"
        "2. שיפור פרגמטיקה חברתית\n"
        "3. תרגול שטף דיבור"
    )

    # Row 3
    row3 = table.rows[3].cells
    row3[0].text = "01/03/2025"
    row3[1].text = (
        "1. הכנה לקראת בדיקות בית ספר\n"
        "2. שיפור הבעת שפה בהקשר כיתתי\n"
        "3. תמיכה הורית ועבודה בבית"
    )

    doc.add_paragraph()

    # --- Layer 2: Session summaries (5 blocks) ---
    _heading(doc, "סיכומי פגישות")
    doc.add_paragraph()

    sessions = [
        {
            "date": "07/01/2025",
            "text": (
                "עבדנו על תרגילי זיכרון שמיעתי — סדרות של מספרים בסדרות עולות. "
                "התגובה הייתה טובה; הצליח לחזור על 4 פריטים. "
                "שיעורי בית: תרגול יומי עם ההורים."
            ),
        },
        {
            "date": "14/01/2025",
            "text": (
                "עבדנו על מודעות פונולוגית — חרוזים וניתוח הברות. "
                "התגובה הייתה טובה מאוד; שיפור ברור מהפגישה הקודמת. "
                "ממשיך לתרגל בבית."
            ),
        },
        {
            "date": "21/01/2025",
            "text": (
                "עבדנו על הרחבת אוצר מילים בנושא בית הספר. "
                "התגובה הייתה ממוצעת; קושי עם מילים מופשטות. "
                "לחזור על החומר בשבוע הבא."
            ),
        },
        {
            "date": "04/02/2025",
            "text": (
                "עבדנו על פרגמטיקה — שיחות חברתיות ותורות דיבור. "
                "התגובה הייתה מצוינת; הפגין יכולת שמירה על שיחה. "
                "מומלץ לתרגל עם חברים."
            ),
        },
        {
            "date": "11/02/2025",
            "text": (
                "עבדנו על שטף דיבור — קריאה בקול. "
                "התגובה הייתה טובה; שטף בסיסי תקין. "
                "לחזור על תרגילי נשימה בבית."
            ),
        },
    ]

    for session in sessions:
        doc.add_paragraph(session["date"])
        doc.add_paragraph(session["text"])
        doc.add_paragraph()

    doc.add_paragraph()
    doc.add_paragraph("___________________________")
    doc.add_paragraph("ד\"ר שרה לוי, קלינאית תקשורת")
    doc.add_paragraph("מ.ר. 12345")

    doc.save(output_path)
    print(f"Created: {output_path}")


# ---------------------------------------------------------------------------
# Entry point
# ---------------------------------------------------------------------------

if __name__ == "__main__":
    create_diagnosis_doc(os.path.join(OUTPUT_DIR, "diagnosis_test.docx"))
    create_clinic_visit_summary_doc(
        os.path.join(OUTPUT_DIR, "clinic_visit_summary_test.docx")
    )
    create_treatment_plan_doc(os.path.join(OUTPUT_DIR, "treatment_plan_test.docx"))
    print("\nAll synthetic test fixtures created successfully.")
